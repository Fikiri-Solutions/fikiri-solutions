"""
AI Document Processing System for Fikiri Solutions
Handles OCR, PDF parsing, text extraction, and document analysis
"""

import os
import json
import logging
import base64
import io
from typing import Dict, Any, List, Optional, Union
from dataclasses import dataclass
from datetime import datetime
from pathlib import Path
import mimetypes

# Optional dependencies with fallbacks
try:
    import pytesseract
    from PIL import Image
    OCR_AVAILABLE = True
except ImportError:
    OCR_AVAILABLE = False

try:
    import PyPDF2
    PDF_AVAILABLE = True
except ImportError:
    PDF_AVAILABLE = False

try:
    import docx
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

try:
    import openpyxl
    EXCEL_AVAILABLE = True
except ImportError:
    EXCEL_AVAILABLE = False

from core.minimal_config import get_config

logger = logging.getLogger(__name__)

@dataclass
class DocumentMetadata:
    """Document metadata structure"""
    filename: str
    file_type: str
    file_size: int
    pages: int
    language: str
    created_at: datetime
    processed_at: datetime
    processing_time: float

@dataclass
class ExtractedContent:
    """Extracted content from document"""
    text: str
    structured_data: Dict[str, Any]
    entities: List[Dict[str, Any]]
    confidence: float
    metadata: DocumentMetadata

@dataclass
class ProcessingResult:
    """Document processing result"""
    success: bool
    content: Optional[ExtractedContent]
    error: Optional[str]
    processing_time: float

class AIDocumentProcessor:
    """AI-powered document processing system"""
    
    def __init__(self):
        self.config = get_config()
        self.supported_formats = {
            'image': ['.jpg', '.jpeg', '.png', '.gif', '.bmp', '.tiff'],
            'pdf': ['.pdf'],
            'document': ['.docx', '.doc'],
            'spreadsheet': ['.xlsx', '.xls', '.csv'],
            'text': ['.txt', '.md', '.rtf']
        }
        
        # Initialize OCR if available
        if OCR_AVAILABLE:
            try:
                # Set Tesseract path if configured
                tesseract_path = getattr(self.config, 'tesseract_path', None)
                if tesseract_path:
                    pytesseract.pytesseract.tesseract_cmd = tesseract_path
                logger.info("✅ OCR (Tesseract) initialized successfully")
            except Exception as e:
                logger.warning(f"⚠️ OCR initialization failed: {e}")
        
        logger.info(f"📄 Document processor initialized - OCR: {OCR_AVAILABLE}, PDF: {PDF_AVAILABLE}, DOCX: {DOCX_AVAILABLE}, Excel: {EXCEL_AVAILABLE}")
    
    def process_document(self, file_path: str, file_content: Optional[bytes] = None) -> ProcessingResult:
        """Process a document and extract content"""
        start_time = datetime.now()
        
        try:
            # Determine file type
            file_type = self._get_file_type(file_path)
            
            if file_type not in self.supported_formats:
                return ProcessingResult(
                    success=False,
                    content=None,
                    error=f"Unsupported file type: {file_type}",
                    processing_time=0
                )
            
            # Extract content based on file type
            if file_type == 'image':
                content = self._process_image(file_path, file_content)
            elif file_type == 'pdf':
                content = self._process_pdf(file_path, file_content)
            elif file_type == 'document':
                content = self._process_document(file_path, file_content)
            elif file_type == 'spreadsheet':
                content = self._process_spreadsheet(file_path, file_content)
            elif file_type == 'text':
                content = self._process_text(file_path, file_content)
            else:
                return ProcessingResult(
                    success=False,
                    content=None,
                    error=f"Unknown file type: {file_type}",
                    processing_time=0
                )
            
            processing_time = (datetime.now() - start_time).total_seconds()
            
            return ProcessingResult(
                success=True,
                content=content,
                error=None,
                processing_time=processing_time
            )
            
        except Exception as e:
            processing_time = (datetime.now() - start_time).total_seconds()
            logger.error(f"❌ Document processing failed: {e}")
            return ProcessingResult(
                success=False,
                content=None,
                error=str(e),
                processing_time=processing_time
            )
    
    def _get_file_type(self, file_path: str) -> str:
        """Determine file type from extension"""
        ext = Path(file_path).suffix.lower()
        
        for file_type, extensions in self.supported_formats.items():
            if ext in extensions:
                return file_type
        
        return 'unknown'
    
    def _process_image(self, file_path: str, file_content: Optional[bytes] = None) -> ExtractedContent:
        """Process image files using OCR"""
        if not OCR_AVAILABLE:
            return self._fallback_image_processing(file_path, file_content)
        
        try:
            # Load image
            if file_content:
                image = Image.open(io.BytesIO(file_content))
            else:
                image = Image.open(file_path)
            
            # Extract text using OCR
            text = pytesseract.image_to_string(image, lang='eng')
            
            # Get additional metadata
            data = pytesseract.image_to_data(image, output_type=pytesseract.Output.DICT)
            
            # Calculate confidence
            confidences = [int(conf) for conf in data['conf'] if int(conf) > 0]
            avg_confidence = sum(confidences) / len(confidences) if confidences else 0
            
            # Extract entities (basic)
            entities = self._extract_entities_from_text(text)
            
            # Create metadata
            metadata = DocumentMetadata(
                filename=Path(file_path).name,
                file_type='image',
                file_size=len(file_content) if file_content else os.path.getsize(file_path),
                pages=1,
                language='en',
                created_at=datetime.now(),
                processed_at=datetime.now(),
                processing_time=0
            )
            
            return ExtractedContent(
                text=text.strip(),
                structured_data={'ocr_data': data},
                entities=entities,
                confidence=avg_confidence / 100.0,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"❌ Image processing failed: {e}")
            return self._fallback_image_processing(file_path, file_content)
    
    def _process_pdf(self, file_path: str, file_content: Optional[bytes] = None) -> ExtractedContent:
        """Process PDF files"""
        if not PDF_AVAILABLE:
            return self._fallback_pdf_processing(file_path, file_content)
        
        try:
            text = ""
            pages = 0
            
            if file_content:
                pdf_reader = PyPDF2.PdfReader(io.BytesIO(file_content))
            else:
                pdf_reader = PyPDF2.PdfReader(file_path)
            
            pages = len(pdf_reader.pages)
            
            for page in pdf_reader.pages:
                text += page.extract_text() + "\n"
            
            # Extract entities
            entities = self._extract_entities_from_text(text)
            
            # Create metadata
            metadata = DocumentMetadata(
                filename=Path(file_path).name,
                file_type='pdf',
                file_size=len(file_content) if file_content else os.path.getsize(file_path),
                pages=pages,
                language='en',
                created_at=datetime.now(),
                processed_at=datetime.now(),
                processing_time=0
            )
            
            return ExtractedContent(
                text=text.strip(),
                structured_data={'pages': pages},
                entities=entities,
                confidence=0.8,  # PDF extraction is generally reliable
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"❌ PDF processing failed: {e}")
            return self._fallback_pdf_processing(file_path, file_content)
    
    def _process_document(self, file_path: str, file_content: Optional[bytes] = None) -> ExtractedContent:
        """Process Word documents"""
        if not DOCX_AVAILABLE:
            return self._fallback_document_processing(file_path, file_content)
        
        try:
            if file_content:
                doc = docx.Document(io.BytesIO(file_content))
            else:
                doc = docx.Document(file_path)
            
            text = ""
            for paragraph in doc.paragraphs:
                text += paragraph.text + "\n"
            
            # Extract tables
            tables_data = []
            for table in doc.tables:
                table_data = []
                for row in table.rows:
                    row_data = [cell.text for cell in row.cells]
                    table_data.append(row_data)
                tables_data.append(table_data)
            
            # Extract entities
            entities = self._extract_entities_from_text(text)
            
            # Create metadata
            metadata = DocumentMetadata(
                filename=Path(file_path).name,
                file_type='document',
                file_size=len(file_content) if file_content else os.path.getsize(file_path),
                pages=len(doc.paragraphs) // 20,  # Rough estimate
                language='en',
                created_at=datetime.now(),
                processed_at=datetime.now(),
                processing_time=0
            )
            
            return ExtractedContent(
                text=text.strip(),
                structured_data={'tables': tables_data},
                entities=entities,
                confidence=0.9,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"❌ Document processing failed: {e}")
            return self._fallback_document_processing(file_path, file_content)
    
    def _process_spreadsheet(self, file_path: str, file_content: Optional[bytes] = None) -> ExtractedContent:
        """Process Excel spreadsheets"""
        if not EXCEL_AVAILABLE:
            return self._fallback_spreadsheet_processing(file_path, file_content)
        
        try:
            if file_content:
                workbook = openpyxl.load_workbook(io.BytesIO(file_content))
            else:
                workbook = openpyxl.load_workbook(file_path)
            
            text = ""
            sheets_data = {}
            
            for sheet_name in workbook.sheetnames:
                sheet = workbook[sheet_name]
                sheet_text = ""
                
                for row in sheet.iter_rows(values_only=True):
                    row_text = " ".join([str(cell) for cell in row if cell is not None])
                    sheet_text += row_text + "\n"
                
                sheets_data[sheet_name] = sheet_text
                text += f"Sheet: {sheet_name}\n{sheet_text}\n"
            
            # Extract entities
            entities = self._extract_entities_from_text(text)
            
            # Create metadata
            metadata = DocumentMetadata(
                filename=Path(file_path).name,
                file_type='spreadsheet',
                file_size=len(file_content) if file_content else os.path.getsize(file_path),
                pages=len(workbook.sheetnames),
                language='en',
                created_at=datetime.now(),
                processed_at=datetime.now(),
                processing_time=0
            )
            
            return ExtractedContent(
                text=text.strip(),
                structured_data={'sheets': sheets_data},
                entities=entities,
                confidence=0.95,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"❌ Spreadsheet processing failed: {e}")
            return self._fallback_spreadsheet_processing(file_path, file_content)
    
    def _process_text(self, file_path: str, file_content: Optional[bytes] = None) -> ExtractedContent:
        """Process plain text files"""
        try:
            if file_content:
                text = file_content.decode('utf-8', errors='ignore')
            else:
                with open(file_path, 'r', encoding='utf-8', errors='ignore') as f:
                    text = f.read()
            
            # Extract entities
            entities = self._extract_entities_from_text(text)
            
            # Create metadata
            metadata = DocumentMetadata(
                filename=Path(file_path).name,
                file_type='text',
                file_size=len(file_content) if file_content else os.path.getsize(file_path),
                pages=len(text.split('\n')) // 50,  # Rough estimate
                language='en',
                created_at=datetime.now(),
                processed_at=datetime.now(),
                processing_time=0
            )
            
            return ExtractedContent(
                text=text.strip(),
                structured_data={'lines': len(text.split('\n'))},
                entities=entities,
                confidence=1.0,
                metadata=metadata
            )
            
        except Exception as e:
            logger.error(f"❌ Text processing failed: {e}")
            raise
    
    def _extract_entities_from_text(self, text: str) -> List[Dict[str, Any]]:
        """Extract entities from text using simple patterns"""
        entities = []
        
        # Email addresses
        import re
        emails = re.findall(r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b', text)
        for email in emails:
            entities.append({
                'type': 'email',
                'value': email,
                'confidence': 0.9
            })
        
        # Phone numbers
        phones = re.findall(r'\b\d{3}[-.]?\d{3}[-.]?\d{4}\b', text)
        for phone in phones:
            entities.append({
                'type': 'phone',
                'value': phone,
                'confidence': 0.8
            })
        
        # URLs
        urls = re.findall(r'http[s]?://(?:[a-zA-Z]|[0-9]|[$-_@.&+]|[!*\\(\\),]|(?:%[0-9a-fA-F][0-9a-fA-F]))+', text)
        for url in urls:
            entities.append({
                'type': 'url',
                'value': url,
                'confidence': 0.9
            })
        
        return entities
    
    def _fallback_image_processing(self, file_path: str, file_content: Optional[bytes] = None) -> ExtractedContent:
        """Fallback for image processing when OCR is not available"""
        metadata = DocumentMetadata(
            filename=Path(file_path).name,
            file_type='image',
            file_size=len(file_content) if file_content else os.path.getsize(file_path),
            pages=1,
            language='en',
            created_at=datetime.now(),
            processed_at=datetime.now(),
            processing_time=0
        )
        
        return ExtractedContent(
            text="[Image file - OCR not available]",
            structured_data={'fallback': True},
            entities=[],
            confidence=0.0,
            metadata=metadata
        )
    
    def _fallback_pdf_processing(self, file_path: str, file_content: Optional[bytes] = None) -> ExtractedContent:
        """Fallback for PDF processing when PyPDF2 is not available"""
        metadata = DocumentMetadata(
            filename=Path(file_path).name,
            file_type='pdf',
            file_size=len(file_content) if file_content else os.path.getsize(file_path),
            pages=1,
            language='en',
            created_at=datetime.now(),
            processed_at=datetime.now(),
            processing_time=0
        )
        
        return ExtractedContent(
            text="[PDF file - PDF processing not available]",
            structured_data={'fallback': True},
            entities=[],
            confidence=0.0,
            metadata=metadata
        )
    
    def _fallback_document_processing(self, file_path: str, file_content: Optional[bytes] = None) -> ExtractedContent:
        """Fallback for document processing when docx is not available"""
        metadata = DocumentMetadata(
            filename=Path(file_path).name,
            file_type='document',
            file_size=len(file_content) if file_content else os.path.getsize(file_path),
            pages=1,
            language='en',
            created_at=datetime.now(),
            processed_at=datetime.now(),
            processing_time=0
        )
        
        return ExtractedContent(
            text="[Document file - Document processing not available]",
            structured_data={'fallback': True},
            entities=[],
            confidence=0.0,
            metadata=metadata
        )
    
    def _fallback_spreadsheet_processing(self, file_path: str, file_content: Optional[bytes] = None) -> ExtractedContent:
        """Fallback for spreadsheet processing when openpyxl is not available"""
        metadata = DocumentMetadata(
            filename=Path(file_path).name,
            file_type='spreadsheet',
            file_size=len(file_content) if file_content else os.path.getsize(file_path),
            pages=1,
            language='en',
            created_at=datetime.now(),
            processed_at=datetime.now(),
            processing_time=0
        )
        
        return ExtractedContent(
            text="[Spreadsheet file - Spreadsheet processing not available]",
            structured_data={'fallback': True},
            entities=[],
            confidence=0.0,
            metadata=metadata
        )
    
    def get_supported_formats(self) -> Dict[str, List[str]]:
        """Get list of supported file formats"""
        return self.supported_formats
    
    def is_format_supported(self, file_path: str) -> bool:
        """Check if file format is supported"""
        file_type = self._get_file_type(file_path)
        return file_type != 'unknown'
    
    def get_processing_capabilities(self) -> Dict[str, bool]:
        """Get processing capabilities status"""
        return {
            'ocr': OCR_AVAILABLE,
            'pdf': PDF_AVAILABLE,
            'docx': DOCX_AVAILABLE,
            'excel': EXCEL_AVAILABLE
        }

# Global instance
document_processor = AIDocumentProcessor()

def get_document_processor() -> AIDocumentProcessor:
    """Get the global document processor instance"""
    return document_processor
